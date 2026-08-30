import streamlit as st
import pandas as pd

from io import BytesIO
from docx import Document

from pdf_reader import extract_text_from_pdf
from docx_reader import extract_text_from_docx
from extractor import analyze_document


def format_nombre(valeur):

    try:

        if valeur is None:
            return "N/A"

        return f"{int(valeur):,}".replace(",", " ")

    except Exception:

        return str(valeur)


def generate_word_report(result):

    doc = Document()

    doc.add_heading(
        "NOTE DE SYNTHESE - SUIVI FONDS",
        level=1
    )

    infos = result.get(
        "informations_generales",
        {}
    )

    performance = result.get(
        "performance",
        {}
    )

    gouvernance = result.get(
        "gouvernance",
        {}
    )

    synthese = result.get(
        "synthese_executive",
        ""
    )

    if synthese:

        doc.add_heading(
            "Synthèse Exécutive",
            level=2
        )

        doc.add_paragraph(
            synthese
        )

    doc.add_heading(
        "Informations Générales",
        level=2
    )

    for cle, valeur in infos.items():

        if cle != "actionnariat":

            doc.add_paragraph(
                f"{cle} : {valeur}"
            )

    doc.add_heading(
        "Performance",
        level=2
    )

    for cle, valeur in performance.items():

        doc.add_paragraph(
            f"{cle} : {valeur}"
        )

    alertes = result.get(
        "alertes",
        []
    )

    if alertes:

        doc.add_heading(
            "Alertes",
            level=2
        )

        for item in alertes:

            doc.add_paragraph(
                f"• {item}"
            )

    risques = result.get(
        "risques",
        []
    )

    if risques:

        doc.add_heading(
            "Risques",
            level=2
        )

        for item in risques:

            doc.add_paragraph(
                f"• {item}"
            )

    decisions = result.get(
        "decisions",
        []
    )

    if decisions:

        doc.add_heading(
            "Décisions",
            level=2
        )

        for item in decisions:

            doc.add_paragraph(
                f"• {item}"
            )

    actions = result.get(
        "actions_a_mener",
        []
    )

    if actions:

        doc.add_heading(
            "Actions de Suivi",
            level=2
        )

        for item in actions:

            doc.add_paragraph(
                f"• {item}"
            )

    doc.add_heading(
        "Gouvernance",
        level=2
    )

    for cle, valeur in gouvernance.items():

        doc.add_paragraph(
            f"{cle} : {valeur}"
        )

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer


st.set_page_config(
    page_title="Suivi Fonds PE et OPCI",
    page_icon="📊",
    layout="wide"
)

st.title("📊 Suivi Fonds PE et OPCI")

uploaded_file = st.file_uploader(
    "Déposer un reporting",
    type=["pdf", "docx"]
)

if uploaded_file:

    if uploaded_file.name.lower().endswith(".pdf"):

        text = extract_text_from_pdf(
            uploaded_file
        )

    else:

        text = extract_text_from_docx(
            uploaded_file
        )

    st.success(
        f"Texte extrait : {len(text)} caractères"
    )

    with st.spinner(
        "Analyse du document..."
    ):

        result = analyze_document(
            text
        )

    if "erreur" in result:

        st.error(
            result["erreur"]
        )

        st.json(result)

        st.stop()

    infos = result.get(
        "informations_generales",
        {}
    )

    perf = result.get(
        "performance",
        {}
    )

    gouvernance = result.get(
        "gouvernance",
        {}
    )

    st.success(
        "Analyse terminée"
    )

    st.header(
        "📌 Synthèse Exécutive"
    )

    synthese = result.get(
        "synthese_executive",
        ""
    )

    if synthese:

        st.info(
            synthese
        )

    col1, col2, col3, col4 = st.columns(4)

    col1.metric(
        "Fonds",
        infos.get(
            "nom_du_fonds",
            "N/A"
        )
    )

    col2.metric(
        "Taille du Fonds",
        infos.get(
            "taille_totale_du_fonds",
            "N/A"
        )
    )

    col3.metric(
        "TRI",
        perf.get(
            "tri",
            "N/A"
        )
    )

    col4.metric(
        "Date Comité",
        gouvernance.get(
            "date_reunion_comite_surveillance",
            "N/A"
        )
    )

    if result.get("alertes"):

        st.header(
            "🚨 Alertes"
        )

        for item in result["alertes"]:

            st.warning(item)

    if result.get("risques"):

        st.header(
            "⚠️ Risques"
        )

        for item in result["risques"]:

            st.error(item)

    if result.get("decisions"):

        st.header(
            "✅ Décisions"
        )

        for item in result["decisions"]:

            st.success(item)

    if result.get("actions_a_mener"):

        st.header(
            "📋 Actions de Suivi"
        )

        for item in result["actions_a_mener"]:

            st.info(item)

    st.header(
        "🏦 Informations Générales"
    )

    st.write(
        f"**Nom du Fonds :** {infos.get('nom_du_fonds','N/A')}"
    )

    st.write(
        f"**Forme juridique :** {infos.get('forme_juridique','N/A')}"
    )

    st.write(
        f"**Nature des investissements :** {infos.get('nature_des_investissements','N/A')}"
    )

    st.write(
        f"**Date de constitution :** {infos.get('date_de_constitution','N/A')}"
    )

    st.write(
        f"**Durée :** {infos.get('duree_du_fonds','N/A')}"
    )

    actionnariat = infos.get(
        "actionnariat",
        {}
    )

    if actionnariat:

        st.header(
            "👥 Actionnariat"
        )

        df_actionnariat = pd.DataFrame(
            list(actionnariat.items()),
            columns=[
                "Investisseur",
                "Participation"
            ]
        )

        st.dataframe(
            df_actionnariat,
            use_container_width=True
        )

    st.header(
        "📈 Performance"
    )

    c1, c2, c3 = st.columns(3)

    c1.metric(
        "Montant Investi",
        format_nombre(
            perf.get(
                "total_investi_dh"
            )
        )
    )

    c1.metric(
        "TRI",
        perf.get(
            "tri",
            "N/A"
        )
    )

    c2.metric(
        "Valorisation",
        format_nombre(
            perf.get(
                "valorisation_portefeuille_dh"
            )
        )
    )

    c2.metric(
        "Valeur liquidative",
        perf.get(
            "valeur_liquidative_actuelle_dh",
            "N/A"
        )
    )

    c3.metric(
        "Plus-value",
        format_nombre(
            perf.get(
                "plus_value_totale_dh"
            )
        )
    )

    c3.metric(
        "Progression VL",
        perf.get(
            "progression_valeur_liquidative",
            "N/A"
        )
    )

    participations = result.get(
        "participations",
        []
    )

    if participations:

        st.header(
            "🏢 Portefeuille de Participations"
        )

        st.dataframe(
            pd.DataFrame(
                participations
            ),
            use_container_width=True
        )

    st.header(
        "🏛 Gouvernance"
    )

    for cle, valeur in gouvernance.items():

        st.write(
            f"**{cle} :** {valeur}"
        )

    st.header(
        "📥 Export"
    )

    rapport_word = generate_word_report(
        result
    )

    st.download_button(
        label="📄 Télécharger la note Word",
        data=rapport_word,
        file_name="Note_Synthese_Fonds.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    with st.expander(
        "Afficher le JSON complet"
    ):

        st.json(result)
